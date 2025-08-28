from flask import Flask, render_template, request, send_file, jsonify
from io import BytesIO
from datetime import datetime
import random
import re

# -------- PDF (reportlab)
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.units import cm


# -------- Word (python-docx)
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


app = Flask(__name__)

@app.route('/googlef9aea8b2c7991914.html')
def google_verification():
    return "google-site-verification: googlef9aea8b2c7991914.html"

import random, re

def humanizar_texto(texto):
    import random, re

    # Diccionario de reemplazos más variado
    reemplazos = {
        "impacto positivo": ["efecto beneficioso", "resultado favorable", "aportación valiosa"],
        "contribuyendo": ["lo cual ayuda a", "favoreciendo", "aportando a"],
        "promoviendo": ["fomentando", "impulsando", "apoyando"],
        "utilizar": ["emplear", "hacer uso de", "usar"],
        "tecnología": ["sistema", "método", "herramienta tecnológica"],
        "optimizar": ["mejorar", "hacer más eficiente", "perfeccionar"],
        "huella de carbono": ["emisiones de CO₂", "rastro ambiental", "contaminación generada"],
        "nuestro": ["la presente", "esta propuesta", "la iniciativa"],
        "ODS": ["Objetivo de Desarrollo Sostenible", "ODS número", "meta de desarrollo sostenible"],
    }

    for original, opciones in reemplazos.items():
        if original in texto:
            texto = texto.replace(original, random.choice(opciones))

    # Cortar frases largas para dar ritmo humano
    frases = re.split(r'(?<=[.!?]) +', texto)
    frases_mod = []
    for f in frases:
        palabras = f.split()
        if len(palabras) > 22:
            corte = random.randint(10, len(palabras) - 8)
            f1 = " ".join(palabras[:corte]) + "."
            f2 = " ".join(palabras[corte:])
            frases_mod.extend([f1, f2])
        else:
            frases_mod.append(f)

    texto = " ".join(frases_mod)

    # Variar inicios de párrafos
    inicios = ["Cabe señalar que", "Es claro que", "En este contexto", "Vale la pena mencionar que"]
    texto = re.sub(r'(^|\n)([A-ZÁÉÍÓÚ])', lambda m: m.group(1) + random.choice(inicios) + " " + m.group(2).lower(), texto, count=2)

    return texto.strip()

# -------------------------------
# RUTAS HUMANIZADOR
# -------------------------------

    
@app.route("/humanizador", methods=["GET", "POST"])
def index():
    resultado = ""
    if request.method == "POST":
        texto = request.form["texto"]
        resultado = humanizar_texto(texto)
    return render_template("index.html", resultado=resultado)

@app.route("/", methods=["GET", "POST"])
def principal():
    resultado = ""
    if request.method == "POST":
        texto = request.form["texto"]
        resultado = humanizar_texto(texto)
    return render_template("principal.html", resultado=resultado)


# -------------------------------
# RUTA GENERADOR (APA completo)
# -------------------------------
@app.route("/generador", methods=["GET", "POST"])
def generador():
    resultado = ""
    if request.method == "POST":
        titulo = request.form.get("titulo", "")
        nombre = request.form.get("nombre", "")
        curso = request.form.get("curso", "")
        profesor = request.form.get("profesor", "")
        institucion = request.form.get("institucion", "")
        facultad = request.form.get("facultad", "")
        ciudad = request.form.get("ciudad", "")
        anio = request.form.get("anio", datetime.now().strftime("%Y"))
        texto = request.form.get("texto", "")
        referencias = request.form.get("referencias", "")
        formato = request.form.get("formato", "pdf")

        # Humanizamos
        resultado = humanizar_texto(texto)

        # Procesamos referencias
        lista_referencias = [ref.strip() for ref in referencias.split("\n") if ref.strip()]

        # ===== PDF =====
        if formato == "pdf":
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=72)

            estilo_texto = ParagraphStyle("APA-Texto",
                                          fontName="Times-Roman",
                                          fontSize=12,
                                          leading=24,
                                          alignment=TA_LEFT,
                                          firstLineIndent=36)

            estilo_centrado = ParagraphStyle("APA-Centrado",
                                             fontName="Times-Roman",
                                             fontSize=12,
                                             leading=24,
                                             alignment=TA_CENTER)

            story = []

            # ---- Portada ----
            story.append(Spacer(1, 200))
            story.append(Paragraph(titulo, estilo_centrado))
            story.append(Spacer(1, 100))
            story.append(Paragraph(nombre, estilo_centrado))
            story.append(Spacer(1, 24))
            story.append(Paragraph(curso, estilo_centrado))
            story.append(Paragraph(profesor, estilo_centrado))
            story.append(Spacer(1, 150))
            story.append(Paragraph(institucion, estilo_centrado))
            story.append(Paragraph(facultad, estilo_centrado))
            story.append(Paragraph(ciudad, estilo_centrado))
            story.append(Paragraph(anio, estilo_centrado))
            story.append(PageBreak())

            # ---- Cuerpo ----
            story.append(Paragraph(titulo, estilo_centrado))  # título del usuario
            for parrafo in resultado.split("\n"):
                if parrafo.strip():
                    story.append(Paragraph(parrafo, estilo_texto))

            # ---- Referencias ----
            if lista_referencias:
                story.append(PageBreak())
                story.append(Paragraph("Referencias", estilo_centrado))
                story.append(Spacer(1, 24))
                for ref in lista_referencias:
                    story.append(Paragraph(ref, estilo_texto))

            doc.build(story)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True,
                             download_name="documento_APA.pdf",
                             mimetype="application/pdf")

        # ===== WORD =====
        elif formato == "word":
            doc = Document()
            section = doc.sections[0]
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)

            # ---- Portada ----
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(titulo).bold = True

            for _ in range(5): doc.add_paragraph()

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(nombre)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(curso)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(profesor)

            for _ in range(8): doc.add_paragraph()

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(institucion)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(facultad)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(ciudad)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(anio)

            doc.add_page_break()

            # ---- Cuerpo ----
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.add_run(titulo).bold = True

            for parrafo in resultado.split("\n"):
                if parrafo.strip():
                    p = doc.add_paragraph(parrafo.strip())
                    pf = p.paragraph_format
                    pf.first_line_indent = Cm(1.27)
                    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # ---- Referencias ----
            if lista_referencias:
                doc.add_page_break()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("Referencias").bold = True

                for ref in lista_referencias:
                    p = doc.add_paragraph(ref)
                    pf = p.paragraph_format
                    pf.first_line_indent = Cm(0)  # APA: sangría francesa opcional
                    pf.left_indent = Cm(1.27)
                    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True,
                             download_name="documento_APA.docx",
                             mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    return render_template("generador.html", resultado=resultado, anio=datetime.now().year)


# -------------------------------
# RUTAS DETECTOR DE IA
# -------------------------------
@app.route("/detector")
def detector():
    return render_template("detector.html")

@app.route("/analizar", methods=["POST"])
def analizar():
    data = request.get_json()
    texto = data.get("texto", "")

    import re, math
    palabras = re.findall(r"[A-Za-zÁÉÍÓÚáéíóúÜüÑñ]+", texto.lower())
    n = len(palabras)
    if n == 0:
        return jsonify({"score": 0})

    # Oraciones y longitudes
    oraciones = [s.strip() for s in re.split(r'[.!?…]+', texto) if s.strip()]
    longitudes = [len(re.findall(r"[A-Za-zÁÉÍÓÚáéíóúÜüÑñ]+", s)) for s in oraciones] or [n]
    media = sum(longitudes) / len(longitudes)
    var = sum((x - media) ** 2 for x in longitudes) / len(longitudes) if len(longitudes) > 1 else 0.0
    std = var ** 0.5
    cv = std / max(media, 1e-9)

    # --- Señales IA ---
    # 1. Burstiness inversa (oraciones muy uniformes => IA)
    burst = 1 - min(cv / 0.8, 1.0)

    # 2. Conectores formales
    conectores = ["además","por lo tanto","en conclusión","sin embargo","por ende","como resultado"]
    texto_lc = texto.lower()
    cnt_conect = sum(texto_lc.count(c) for c in conectores)
    conn = min((cnt_conect / max(len(oraciones), 1)) / 0.5, 1.0)

    # 3. Palabras largas (≥12 letras)
    ratio_largas = sum(1 for w in palabras if len(w) >= 12) / n
    largas = min(ratio_largas / 0.08, 1.0)

    # 4. Perplejidad proxy (diversidad)
    unicas = len(set(palabras))
    diversidad = unicas / n
    perplejidad = 1 - min(max((diversidad - 0.45) / 0.3, 0), 1)  # baja diversidad => más IA

    # --- Señales Humanas ---
    personales = {"yo","nosotros","me","mío","mía","nuestro","oye","jaja","mmm","pues","ok"}
    cnt_personal = sum(1 for w in palabras if w in personales)
    personal = min(cnt_personal / (len(oraciones) or 1), 1.0)

    tipos_punt = sum(1 for ch in "!,;:¿?()" if ch in texto)
    variedad_punt = min(tipos_punt / 4.0, 1.0)

    # --- Combinar ---
    ia_score = 0.35*burst + 0.25*conn + 0.2*largas + 0.2*perplejidad
    human_score = 0.5*personal + 0.5*variedad_punt

    # IA final = señalIA - señalHumana
    final = ia_score - human_score

    # Escalar fuerte a extremos
    if final >= 0.3:
        score = 100
    elif final <= -0.3:
        score = 0
    else:
        score = int((final + 0.3) / 0.6 * 100)

    return jsonify({"score": score})


if __name__ == "__main__":
    app.run(debug=True, port= 5001)
